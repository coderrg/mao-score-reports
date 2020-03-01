# Contact Rhythm Garg (typhoonrg@gmail.com) with questions

import pandas as pd
import os
import shutil
from docx import Document
from docx.shared import Inches
from Levenshtein import distance as levenshtein_distance

# define year-specific things
year = 2020
amc12_a_cutoff = 87
amc12_b_cutoff = 87
fma_a_cutoff = 16
fma_b_cutoff = 16

# Read in AMC scores
amc12a_scores = pd.read_csv(str(year) + "AMC12A-TAMS.csv")
# amc12b_scores = pd.read_csv("2020AMC12B-TAMS.csv")


# Read in F=ma scores
fma_scores = pd.read_csv(str(year) + "Fma-TAMS.csv")

# Dictionary for storing all scores
names_and_scores = {}

# Add AMC12A scores to dictionary
for index, row in amc12a_scores.iterrows():
    name = (row["FirstName"] + " " + row["LastName"]).upper()
    score = float(row["TotalScore"])
    if (name in names_and_scores):
        names_and_scores[name].append("AMC12A Score: " + str(score))
    else:
        names_and_scores[name] = ["AMC12A Score: " + str(score)]

'''
# Add AMC12B scores to dictionary
for index, row in amc12b_scores.iterrows():
    name = (row["FirstName"] +  " " + row["LastName"]).upper()
    score = float(row["TotalScore"])
    if (name in names_and_scores):
        names_and_scores[name].append("AMC12B Score: " + str(score))
    else:
        names_and_scores[name] = ["AMC12B Score: " + str(score)]
'''

# Add F=ma scores to dictionary
for index, row in fma_scores.iterrows():
    name = (row["Name"].split()[0]+ " " + row["Name"].split()[-1]).upper()
    if (row["Fma A Score"] >= 0):
        score = int(row["Fma A Score"])
        if (name in names_and_scores):
            names_and_scores[name].append("F=ma A Score: " + str(score))
        else:
            names_and_scores[name] = ["F=ma A Score: " + str(score)]
    else:
        score = int(row["Fma B Score"])
        if (name in names_and_scores):
            names_and_scores[name].append("F=ma B Score: " + str(score))
        else:
            names_and_scores[name] = ["F=ma B Score: " + str(score)]

print(len(names_and_scores))

# Combine names that are close together (account for scantron misbubbles on name)
new_names_and_scores = {}
covered = []
for student, scores in names_and_scores.items():
    if (student not in covered):
        covered.append(student)
        studentName = student
        studentScores = scores
        for student2, scores2 in names_and_scores.items():
            if (levenshtein_distance(studentName, student2) > 0):
                if (levenshtein_distance(studentName, student2) < 3):
                    # choose the name that was used more often
                    if (len(names_and_scores[studentName]) >= len(names_and_scores[student2])):
                        print("Misbubble: " + student2)
                        studentScores.append(names_and_scores[student2])
                    else:
                        print("Misbubble: " + studentName)
                        studentName = student2
                        studentScores.append(names_and_scores[student2])
                    covered.append(student2)
        new_names_and_scores[studentName] = studentScores

print(len(new_names_and_scores))

shutil.rmtree("student_scorereports")
os.mkdir("student_scorereports")
os.chdir("student_scorereports")

for student, scores in new_names_and_scores.items():
    print(student)
    print(scores)

    document = Document()

    document.add_heading("Thank you for competing with TAMS Mu Alpha Theta, " + student + "!", 0)

    document.add_paragraph("Here are the cutoffs this year:")

    document.add_paragraph(str(year) + " AMC12A – " + str(amc12_a_cutoff) + " for AIME Qualification", style='List Bullet')
    
    document.add_paragraph(str(year) + " AMC12B – " + str(amc12_b_cutoff) + " for AIME Qualification", style='List Bullet')
    
    document.add_paragraph(str(year) + " F=ma A – " + str(fma_a_cutoff) + " for USAPhO Qualification", style='List Bullet')
    
    document.add_paragraph(str(year) + " F=ma B – " + str(fma_b_cutoff) + " for USAPhO Qualification", style='List Bullet')

    document.add_paragraph("Below are your contest scores on AMC/F=ma for the 2019-2020 school year with TAMS. Please message a MAO exec if you have any questions.")
    
    for score in scores:
        document.add_paragraph(score)

    document.save(str(year) + " " + str(student) + " TAMS MAO Contest Scores" + ".docx")

